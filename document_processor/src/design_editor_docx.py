from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_page_color_docx(doc_path: str, output_path: str, hex_color: str):
    """
    Sets the page background color for a DOCX document using OOXML manipulation.
    hex_color should be a 6-digit hex string (e.g., "FFFF00" for yellow).
    Saves the modified document to output_path.
    """
    try:
        doc = Document(doc_path)
        
        # Remove 'FF' from alpha if present, Word uses 6-digit hex
        if len(hex_color) == 8 and hex_color.startswith("FF"): # Check if it's alpha FF
            hex_color = hex_color[2:]
        elif len(hex_color) == 8: # If alpha is not FF, Word might not support it well this way.
            print(f"Warning: Hex color {hex_color} includes alpha; Word page color typically uses 6-digit RRGGBB. Proceeding with last 6 digits.")
            hex_color = hex_color[2:]

        if len(hex_color) != 6:
            print(f"Invalid hex color format: {hex_color}. Must be 6-digit (e.g., RRGGBB).")
            return False

        # Access the settings part
        settings_part = doc.settings.element # doc.part.settings_part (python-docx < 0.11.0) or doc.settings.element (>=0.11.0)

        # Find or create the w:background element
        background_tag = settings_part.find(qn('w:background'))
        if background_tag is None:
            background_tag = OxmlElement('w:background')
            # Insert background_tag into settings_part.
            # A common place is before elements like w:evenAndOddHeaders, w:mirrorMargins, etc.
            # Or simply append if structure is not critical, though Word might reorder it.
            # For robustness, find a known element and insert before or after, or prepend/append.
            # Here, we try to append it to settings.xml's root children.
            settings_part.append(background_tag)


        # Remove existing color attributes or other background types to ensure our color is applied
        # List of attributes that might define background color/fill in <w:background>
        # Common ones: w:color, w:themeColor, w:themeTint, w:themeShade
        # We want to set w:color, so remove others.
        attrs_to_clear = [
            qn('w:themeColor'), qn('w:themeTint'), qn('w:themeShade'),
            # also clear w:color itself in case it was set to something else.
            qn('w:color') 
        ]
        for attr_qn in attrs_to_clear:
            if background_tag.get(attr_qn) is not None:
                del background_tag.attrib[attr_qn]
        
        # Also remove child elements like w:drawing, if any (though less common for simple color)
        # For this basic implementation, we'll focus on attributes.

        # Set the new color attribute
        background_tag.set(qn('w:color'), hex_color)
        
        # Optional: ensure other attributes that might affect visibility are not set, e.g., w:displayBackgroundShape="0"
        # Forcing display (if Word respects it):
        # background_tag.set(qn('w:displayBackgroundShape'), '1') # Not standard, Word usually shows color if w:color is set.


        doc.save(output_path)
        print(f"DOCX page color set to {hex_color} and saved to {output_path}")
        return True
    except Exception as e:
        print(f"Error in set_page_color_docx: {e}")
        # import traceback
        # traceback.print_exc()
        return False

def set_text_color_docx(doc_path: str, output_path: str, hex_color: str):
    """
    Sets the text color for the entire DOCX document.
    hex_color should be a 6-digit hex string (e.g., "FF0000" for red).
    Saves the modified document to output_path.
    """
    try:
        doc = Document(doc_path)
        if not (len(hex_color) == 6 and all(c in '0123456789abcdefABCDEF' for c in hex_color)):
            print("Invalid hex color. Must be 6 digits (RRGGBB).")
            return False
        
        color = RGBColor.from_string(hex_color)

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = color
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = color
        
        doc.save(output_path)
        print(f"DOCX text color set to {hex_color} and saved to {output_path}")
        return True
    except Exception as e:
        print(f"Error in set_text_color_docx: {e}")
        return False

def set_font_properties_docx(doc_path: str, output_path: str, font_name: str = None, font_size_pt: float = None):
    """
    Sets font name and/or size for the entire DOCX document.
    Saves the modified document to output_path.
    """
    try:
        doc = Document(doc_path)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if font_name:
                    run.font.name = font_name
                    # For complex scripts (e.g. Arabic, Hebrew, East Asian languages)
                    # it's often necessary to set these as well.
                    run.font.cs_name = font_name  # Complex Script font
                    run.font.h_ansi = font_name # High ANSI font (often same as ASCII)
                    # run.font.east_asia_name = font_name # If targeting East Asian languages specifically
                if font_size_pt:
                    run.font.size = Pt(font_size_pt)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if font_name:
                                run.font.name = font_name
                                run.font.cs_name = font_name
                                run.font.h_ansi = font_name
                                # run.font.east_asia_name = font_name
                            if font_size_pt:
                                run.font.size = Pt(font_size_pt)
        
        doc.save(output_path)
        print(f"DOCX font properties (Name: {font_name}, Size: {font_size_pt}pt) set and saved to {output_path}")
        return True
    except Exception as e:
        print(f"Error in set_font_properties_docx: {e}")
        return False

def add_simple_page_numbers_docx(doc_path: str, output_path: str):
    """
    Adds simple page numbers (bottom center) to each section's footer of a DOCX document.
    Saves the modified document to output_path.
    """
    try:
        doc = Document(doc_path)
        for section_idx, section in enumerate(doc.sections):
            # Ensure each section has its own footer, not linked to the previous one
            # This is important if different sections need different footers or no footer
            if section_idx > 0: # For the first section, it cannot be linked to a "previous" one
                 section.footer.is_linked_to_previous = False

            footer = section.footer
            if not footer.paragraphs:
                p = footer.add_paragraph()
            else:
                # If there's content, we might want to clear it or append
                # For simplicity, let's use the first paragraph or create one if empty.
                # A more robust solution might involve finding an empty paragraph or clearing existing ones.
                p = footer.paragraphs[0] 
                # Clear existing content in the paragraph to avoid multiple page numbers if run again
                for run in p.runs:
                    p._element.remove(run._r)
            
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Create the PAGEREF field
            run = p.add_run()
            fldChar_begin = OxmlElement('w:fldChar')
            fldChar_begin.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE' # Field code for current page number
            
            fldChar_separate = OxmlElement('w:fldChar') # Optional: Separator for complex fields
            fldChar_separate.set(qn('w:fldCharType'), 'separate')

            # Run for displaying the page number (optional, Word usually generates this)
            # run_display = p.add_run() 
            # no_break_hyphen = OxmlElement('w:noBreakHyphen') # For display if field is not updated
            # run_display._r.append(no_break_hyphen)


            fldChar_end = OxmlElement('w:fldChar')
            fldChar_end.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar_begin)
            run._r.append(instrText)
            # If Word doesn't automatically show a number, it might need a w:t element within the field
            # or for the separate character to be there.
            # A common structure is begin -> instrText -> separate -> run with w:t (for cached result) -> end
            # However, often just begin -> instrText -> end is enough for Word to populate.
            # Let's try adding the separate and end markers to make it more robust.
            # run._r.append(fldChar_separate) # Not strictly needed for simple PAGE field
            run._r.append(fldChar_end)
            
            # Ensure the paragraph is not empty if it was cleared
        if not p.text and not p.runs:
                 # Add a default run if paragraph is completely empty to ensure footer is visible
                 # This might not be necessary if the field itself makes the paragraph non-empty
                 # p.add_run(" ") # Placeholder if needed, but field should suffice

            doc.save(output_path)
        print(f"DOCX simple page numbers added and saved to {output_path}")
        return True
    except Exception as e:
        print(f"Error in add_simple_page_numbers_docx: {e}")
        import traceback
        traceback.print_exc()
        return False

import re

def bold_keywords_docx(doc_path: str, output_path: str, keywords: list[str]):
    """
    Finds specified keywords in a DOCX document and makes them bold.
    The search is case-insensitive. Keywords should be plain text strings.
    Saves the modified document to output_path.
    """
    if not keywords:
        print("No keywords provided for bolding.")
        import shutil
        try:
            shutil.copy(doc_path, output_path)
        except Exception as e:
            print(f"Error copying document when no keywords provided: {e}")
            return False # Failed to produce output
        return True


    try:
        doc = Document(doc_path)
        # Create a single regex for all keywords for efficiency, case-insensitive
        # Sort keywords by length (descending) to match longer phrases first
        keywords.sort(key=len, reverse=True)
        # Filter out empty keywords if any, as they can cause issues with regex
        valid_keywords = [kw for kw in keywords if kw]
        if not valid_keywords:
            print("No valid (non-empty) keywords provided.")
            import shutil
            shutil.copy(doc_path, output_path) # Save a copy as no operation will be performed
            return True

        keyword_regex = re.compile(r'|'.join(map(re.escape, valid_keywords)), re.IGNORECASE)

        for para in doc.paragraphs:
            if not para.runs: # Skip if paragraph has no runs
                continue

            full_text = para.text
            if not keyword_regex.search(full_text):
                continue 

            current_runs = list(para.runs) 
            para.clear() 

            run_attributes = {} 
            if current_runs:
                original_first_run = current_runs[0]
                run_attributes['name'] = original_first_run.font.name
                run_attributes['size'] = original_first_run.font.size
                run_attributes['italic'] = original_first_run.font.italic
                run_attributes['underline'] = original_first_run.font.underline
                run_attributes['color_rgb'] = original_first_run.font.color.rgb if original_first_run.font.color and hasattr(original_first_run.font.color, 'rgb') else None
            
            last_end = 0
            for match in keyword_regex.finditer(full_text):
                start, end = match.span()
                keyword_found = match.group(0)

                if start > last_end:
                    run = para.add_run(full_text[last_end:start])
                    if run_attributes.get('name'): run.font.name = run_attributes['name']
                    if run_attributes.get('size'): run.font.size = run_attributes['size']
                    if run_attributes.get('italic'): run.font.italic = run_attributes['italic']
                    if run_attributes.get('underline'): run.font.underline = run_attributes['underline']
                    if run_attributes.get('color_rgb'): run.font.color.rgb = run_attributes['color_rgb']

                bold_run = para.add_run(keyword_found)
                bold_run.bold = True
                if run_attributes.get('name'): bold_run.font.name = run_attributes['name']
                if run_attributes.get('size'): bold_run.font.size = run_attributes['size']
                # Typically, bolding overrides italic for the bolded keyword itself, unless specifically desired
                # if run_attributes.get('italic'): bold_run.font.italic = run_attributes['italic'] 
                if run_attributes.get('underline'): bold_run.font.underline = run_attributes['underline'] # Preserve underline if original was underlined
                if run_attributes.get('color_rgb'): bold_run.font.color.rgb = run_attributes['color_rgb']

                last_end = end
            
            if last_end < len(full_text):
                run = para.add_run(full_text[last_end:])
                if run_attributes.get('name'): run.font.name = run_attributes['name']
                if run_attributes.get('size'): run.font.size = run_attributes['size']
                if run_attributes.get('italic'): run.font.italic = run_attributes['italic']
                if run_attributes.get('underline'): run.font.underline = run_attributes['underline']
                if run_attributes.get('color_rgb'): run.font.color.rgb = run_attributes['color_rgb']

        # Repeat for tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.runs: continue
                        full_text = para.text
                        if not keyword_regex.search(full_text): continue
                        
                        current_runs_table = list(para.runs)
                        para.clear()
                        run_attributes_table = {}
                        if current_runs_table:
                            original_first_run_table = current_runs_table[0]
                            run_attributes_table['name'] = original_first_run_table.font.name
                            run_attributes_table['size'] = original_first_run_table.font.size
                            run_attributes_table['italic'] = original_first_run_table.font.italic
                            run_attributes_table['underline'] = original_first_run_table.font.underline
                            run_attributes_table['color_rgb'] = original_first_run_table.font.color.rgb if original_first_run_table.font.color and hasattr(original_first_run_table.font.color, 'rgb') else None

                        last_end = 0
                        for match in keyword_regex.finditer(full_text):
                            start, end = match.span()
                            keyword_found = match.group(0)
                            if start > last_end:
                                run = para.add_run(full_text[last_end:start])
                                if run_attributes_table.get('name'): run.font.name = run_attributes_table['name']
                                if run_attributes_table.get('size'): run.font.size = run_attributes_table['size']
                                if run_attributes_table.get('italic'): run.font.italic = run_attributes_table['italic']
                                if run_attributes_table.get('underline'): run.font.underline = run_attributes_table['underline']
                                if run_attributes_table.get('color_rgb'): run.font.color.rgb = run_attributes_table['color_rgb']

                            bold_run = para.add_run(keyword_found)
                            bold_run.bold = True
                            if run_attributes_table.get('name'): bold_run.font.name = run_attributes_table['name']
                            if run_attributes_table.get('size'): bold_run.font.size = run_attributes_table['size']
                            # if run_attributes_table.get('italic'): bold_run.font.italic = run_attributes_table['italic'] 
                            if run_attributes_table.get('underline'): bold_run.font.underline = run_attributes_table['underline']
                            if run_attributes_table.get('color_rgb'): bold_run.font.color.rgb = run_attributes_table['color_rgb']
                            last_end = end

                        if last_end < len(full_text): # This is the 'if' statement from the error context
                            run = para.add_run(full_text[last_end:]) # This line and subsequent attribute settings must be indented under it.
                            # Start of explicitly 4-space indented block
                            if run_attributes_table.get('name'):
                                run.font.name = run_attributes_table['name']
                            if run_attributes_table.get('size'):
                                run.font.size = run_attributes_table['size']
                            if run_attributes_table.get('italic'):
                                run.font.italic = run_attributes_table['italic']
                            if run_attributes_table.get('underline'):
                                run.font.underline = run_attributes_table['underline']
                            if run_attributes_table.get('color_rgb'):
                                run.font.color.rgb = run_attributes_table['color_rgb']
                            # End of explicitly 4-space indented block

        doc.save(output_path) # This should be outside the loops, correctly placed as per file content
        print(f"DOCX keywords bolded and saved to {output_path}")
        return True
    except Exception as e:
        print(f"Error in bold_keywords_docx: {e}")
        # import traceback
        # traceback.print_exc()
        import shutil
        try:
            # Attempt to copy original to output path to prevent data loss on this path
            shutil.copy(doc_path, output_path) 
            print(f"Copied original file to {output_path} due to error during processing.")
        except Exception as copy_e:
            print(f"Failed to copy original file to {output_path} after error: {copy_e}")
        return False
