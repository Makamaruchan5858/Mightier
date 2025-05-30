import os
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont # For Japanese

# Ensure output directory exists
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "test_files") # Assuming scripts/ is sibling to data/
os.makedirs(OUTPUT_DIR, exist_ok=True)

# --- Text Generation ---
TARGET_CHARS = 300000
# A sample Japanese paragraph (approx 100 chars, including punctuation)
# "吾輩は猫である。名前はまだ無い。どこで生れたかとんと見当がつかぬ。何でも薄暗いじめじめした所でニャーニャー泣いていた事だけは記憶している。" (Natsume Soseki)
# This is 71 characters.
# sample_paragraph = "これは大規模ファイル処理の検証と改善のために使用される、約30万文字の日本語テストドキュメントです。この文章は繰り返し生成され、指定された文字数に達するまで続きます。
" # approx 100 chars
sample_paragraph = "吾輩は猫である。名前はまだ無い。どこで生れたかとんと見当がつかぬ。何でも薄暗いじめじめした所でニャーニャー泣いていた事だけは記憶している。吾輩はここで始めて人間というものを見た。\n" # approx 100 chars

generated_text = ""
while len(generated_text) < TARGET_CHARS:
    generated_text += sample_paragraph

generated_text = generated_text[:TARGET_CHARS] # Trim to exact target length
print(f"Generated text with {len(generated_text)} characters.")

# --- DOCX Generation ---
docx_path = os.path.join(OUTPUT_DIR, "large_test_JA.docx")
try:
    doc = Document()
    # Add text in multiple paragraphs for better structure if desired,
    # or just one large block for simplicity.
    # For simplicity, add as distinct paragraphs to avoid one massive paragraph.
    # Split generated_text into chunks of e.g. 5 sample_paragraph lengths
    # num_para_units = len(generated_text) // len(sample_paragraph) # Not used directly
    single_block_len = len(sample_paragraph) * 5 
    
    current_pos = 0
    while current_pos < len(generated_text):
        end_pos = min(current_pos + single_block_len, len(generated_text)) # Ensure not to go past the end
        doc.add_paragraph(generated_text[current_pos:end_pos])
        current_pos = end_pos
        # No need to add extra newline paragraph if source text already contains \n
        # if current_pos < len(generated_text): # Add extra newline for spacing between blocks unless it's the last one
        #      doc.add_paragraph("") 


    # doc.add_paragraph(generated_text) # Simpler: one giant paragraph
    doc.save(docx_path)
    print(f"Generated DOCX: {docx_path} ({os.path.getsize(docx_path) / (1024*1024):.2f} MB)")
except Exception as e:
    print(f"Error generating DOCX: {e}")

# --- PDF Generation ---
pdf_path = os.path.join(OUTPUT_DIR, "large_test_JA.pdf")
try:
    # Register a Japanese font (ensure the font file is accessible or use standard CID fonts)
    # Using a standard CID font available with ReportLab
    # pdfmetrics.registerFont(UnicodeCIDFont('HeiseiMin-W3'))
    # pdfmetrics.registerFont(UnicodeCIDFont('HeiseiKakuGo-W5'))
    # Let's try HeiseiKakuGo-W5 as it's a common Gothic font
    FONT_NAME = 'HeiseiKakuGo-W5'
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(FONT_NAME))
    except Exception as font_e:
        print(f"Could not register CID font {FONT_NAME}, falling back to STSong-Light (Chinese, might work for some JP): {font_e}")
        FONT_NAME = 'STSong-Light' # A fallback CJK font, may not be ideal for Japanese
        pdfmetrics.registerFont(UnicodeCIDFont(FONT_NAME))


    c = canvas.Canvas(pdf_path, pagesize=A4)
    c.setFont(FONT_NAME, 10) # Set font and size

    text_object = c.beginText()
    text_object.setTextOrigin(15*mm, A4[1] - 20*mm) # Start near top-left
    text_object.setFont(FONT_NAME, 10) # Ensure font is set on text object
    
    line_height_pt = 12 # points (approx 4.2mm for 10pt font)
    margin_top_bottom_mm = 20 # mm
    page_height_mm = A4[1] / mm
    
    # Calculate max lines per page based on usable height and line height
    usable_page_height_mm = page_height_mm - (2 * margin_top_bottom_mm)
    max_lines_per_page = int(usable_page_height_mm / (line_height_pt * mm / (1*mm * pdfmetrics.getFont(FONT_NAME).face.ascent / 1000) ) ) # More accurate using font ascent
    # Simplified approximation: max_lines_per_page = int( (A4[1] - 40*mm) / (line_height_pt * mm / (1*mm) ) )
    # The previous calculation was: int((A4[1] - 40*mm) / (line_height * mm / (1*mm) ) ) which is similar.
    # Let's use a slightly simpler but effective version:
    max_lines_per_page = int( (A4[1]/mm - 2 * margin_top_bottom_mm) / (line_height_pt * 0.352778) ) # 1pt = 0.352778mm


    lines = generated_text.splitlines()
    current_line_on_page = 0

    for line in lines:
        if not line.strip() and current_line_on_page == 0 and text_object.getY() == (A4[1] - 20*mm): # Skip empty leading lines on new page
            continue

        # Check if we need a new page
        # textObject.getY() decreases as lines are added. Check against bottom margin.
        # Or, simpler, count lines like before.
        if current_line_on_page >= max_lines_per_page:
            c.drawText(text_object)
            c.showPage()
            c.setFont(FONT_NAME, 10) # Re-set font for new page
            text_object = c.beginText()
            text_object.setTextOrigin(15*mm, A4[1] - 20*mm)
            text_object.setFont(FONT_NAME, 10) # Ensure font is set
            current_line_on_page = 0
            if not line.strip(): # if the line that triggered page break was empty, skip it on new page
                continue


        text_object.textLine(line)
        current_line_on_page += 1
    
    c.drawText(text_object) # Draw any remaining text
    c.save()
    print(f"Generated PDF: {pdf_path} ({os.path.getsize(pdf_path) / (1024*1024):.2f} MB)")
except Exception as e:
    print(f"Error generating PDF: {e}")
    import traceback
    traceback.print_exc()
